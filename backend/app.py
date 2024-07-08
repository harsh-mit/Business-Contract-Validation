from flask import Flask, request, jsonify, render_template, send_file
from transformers import pipeline
from xhtml2pdf import pisa
import fitz  # PyMuPDF
import os
from io import BytesIO
import logging
from docx import Document

app = Flask(__name__)

# Setup logging
logging.basicConfig(level=logging.DEBUG)

# Load NER model
ner_model = pipeline("ner", grouped_entities=True)

# Define keywords to highlight
highlight_keywords = [
    "agreement", "confidentiality", "party", "termination", "liability",
    "governing law", "jurisdiction", "indemnity", "warranty", "obligation",
    "payment", "force majeure", "amendment", "assignment", "severability",
    "dispute resolution", "arbitration", "intellectual property",
    "limitation of liability", "entire agreement"
]

def convert_html_to_pdf(source_html, output_filename):
    result_file = BytesIO()
    # convert HTML to PDF
    pisa_status = pisa.CreatePDF(
        source_html,                # the HTML to convert
        dest=result_file)           # file handle to receive result
    result_file.seek(0)
    if pisa_status.err:
        return None
    return result_file

def extract_text_from_pdf(pdf_file):
    try:
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
        return text
    except Exception as e:
        app.logger.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_word(docx_file):
    try:
        document = Document(docx_file)
        text = ""
        for para in document.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        app.logger.error(f"Error extracting text from Word file: {e}")
        return None

def convert_html_to_word(html_content):
    try:
        doc = Document()
        doc.add_heading('Highlighted Contract', 0)
        paragraphs = html_content.split("<br>")
        for paragraph in paragraphs:
            p = doc.add_paragraph()
            if "<strong style='background-color: yellow;'>" in paragraph:
                parts = paragraph.split("<strong style='background-color: yellow;'>")
                for part in parts:
                    if "</strong>" in part:
                        text, rest = part.split("</strong>", 1)
                        run = p.add_run(text)
                        run.font.highlight_color = 3  # Yellow highlight
                        p.add_run(rest)
                    else:
                        p.add_run(part)
            else:
                p.add_run(paragraph)
        word_file = BytesIO()
        doc.save(word_file)
        word_file.seek(0)
        return word_file
    except Exception as e:
        app.logger.error(f"Error generating Word document: {e}")
        return None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/highlight', methods=['POST'])
def highlight_contract():
    try:
        contract_text = request.form['contract_text']

        if 'contract_file' in request.files:
            contract_file = request.files['contract_file']
            if contract_file.filename != '':
                if contract_file.filename.endswith('.pdf'):
                    contract_text = extract_text_from_pdf(contract_file)
                elif contract_file.filename.endswith('.docx'):
                    contract_text = extract_text_from_word(contract_file)
                else:
                    contract_text = contract_file.read().decode('utf-8')

        if not contract_text:
            return render_template('index.html', error="Failed to extract text from the contract.")

        highlighted_content = contract_text
        for keyword in highlight_keywords:
            highlighted_content = highlighted_content.replace(keyword, f"<strong style='background-color: yellow;'>{keyword}</strong>")

        detected_entities = ner_model(contract_text)
        entities = [{'word': e['word'], 'label': e['entity_group'], 'score': e['score']} for e in detected_entities]

        return render_template('index.html', highlighted_text=highlighted_content, entities=entities, contract_text=contract_text)
    except Exception as e:
        app.logger.error(f"Error highlighting contract: {e}")
        return render_template('index.html', error="There was an error processing your request.")

@app.route('/download', methods=['POST'])
def download_pdf():
    try:
        html_content = request.form['html_content']
        pdf_file = convert_html_to_pdf(html_content, "highlighted_contract.pdf")
        if pdf_file:
            return send_file(pdf_file, as_attachment=True, download_name='highlighted_contract.pdf')
        else:
            return jsonify({"error": "Failed to generate PDF"}), 500
    except Exception as e:
        app.logger.error(f"Error downloading PDF: {e}")
        return jsonify({"error": "There was an error processing your request."}), 500

@app.route('/download_word', methods=['POST'])
def download_word():
    try:
        html_content = request.form['html_content']
        word_file = convert_html_to_word(html_content)
        if word_file:
            return send_file(word_file, as_attachment=True, download_name='highlighted_contract.docx')
        else:
            return jsonify({"error": "Failed to generate Word document"}), 500
    except Exception as e:
        app.logger.error(f"Error downloading Word document: {e}")
        return jsonify({"error": "There was an error processing your request."}), 500

if __name__ == '__main__':
    app.run(debug=True)
